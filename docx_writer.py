"""Consumes DocOp objects and writes a python-docx Document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from epub2docx.html_walker import (
    DocOp, AddHeading, StartParagraph, AddRun, AddSceneBreak, AddLineBreak,
)
from typing import List


def write_docx(ops: List[DocOp], title: str, output_path: str):
    """Write a list of DocOp operations to a DOCX file.

    Args:
        ops: List of DocOp objects from the walker.
        title: The document/chapter title (used as Heading 1).
        output_path: Path to write the .docx file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Add document title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_para = None

    for op in ops:
        if isinstance(op, AddHeading):
            p = doc.add_heading(op.text, level=op.level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_para = None

        elif isinstance(op, AddSceneBreak):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("* * *")
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            current_para = None

        elif isinstance(op, StartParagraph):
            current_para = doc.add_paragraph()
            _apply_paragraph_style(current_para, op.style)

        elif isinstance(op, AddRun):
            if current_para is None:
                # Orphan run - create a default paragraph
                current_para = doc.add_paragraph()
                _apply_paragraph_style(current_para, "body_text")

            if op.text == "\n":
                current_para.add_run("\n")
            elif op.text:
                run = current_para.add_run(op.text)
                run.font.size = Pt(11)
                if op.italic:
                    run.font.italic = True
                if op.bold:
                    run.font.bold = True

        elif isinstance(op, AddLineBreak):
            if current_para is not None:
                current_para.add_run("\n")

    doc.save(output_path)


def append_to_docx(ops: List[DocOp], output_path: str):
    """Append additional DocOp operations to an existing DOCX file.

    Used when an EPUB splits a chapter across multiple HTML files and
    the converter needs to add the continuation to the already-written
    chapter file.

    Args:
        ops: List of DocOp objects to append.
        output_path: Path to the existing .docx file.
    """
    doc = Document(output_path)

    # Add a scene break to visually separate the continuation
    # (it was one chapter in print, this just marks the join point)
    current_para = None

    for op in ops:
        if isinstance(op, AddHeading):
            p = doc.add_heading(op.text, level=op.level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_para = None

        elif isinstance(op, AddSceneBreak):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("* * *")
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            current_para = None

        elif isinstance(op, StartParagraph):
            current_para = doc.add_paragraph()
            _apply_paragraph_style(current_para, op.style)

        elif isinstance(op, AddRun):
            if current_para is None:
                current_para = doc.add_paragraph()
                _apply_paragraph_style(current_para, "body_text")

            if op.text == "\n":
                current_para.add_run("\n")
            elif op.text:
                run = current_para.add_run(op.text)
                run.font.size = Pt(11)
                if op.italic:
                    run.font.italic = True
                if op.bold:
                    run.font.bold = True

        elif isinstance(op, AddLineBreak):
            if current_para is not None:
                current_para.add_run("\n")

    doc.save(output_path)


def _apply_paragraph_style(para, style: str):
    """Apply formatting to a paragraph based on its semantic style."""
    if style == "body_text":
        para.paragraph_format.first_line_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(2)

    elif style == "opening_text":
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.space_after = Pt(2)

    elif style == "heading":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)

    elif style == "author_heading":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(12)

    elif style == "block_quote":
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(4)

    elif style == "dedication":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(6)

    elif style == "epigraph":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(6)

    else:
        # Unknown style - use body text defaults
        para.paragraph_format.first_line_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(2)
